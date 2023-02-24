import streamlit as st
from docx import Document
from docx.shared import Inches
from docx.enum.style import WD_STYLE_TYPE
import io
import datetime

st.set_page_config(page_title="Rapportgenerator", page_icon="üìò")

with open("styles/main.css") as f:
    st.markdown("<style>{}</style>".format(f.read()), unsafe_allow_html=True) 
    
st.title("Rapportgenerator")

st.header("Innhenting av data")
#st.subheader("Eksempel: OpenAI API")
#command = st.text_input("Lag et bilde med AI (skriv inn tekst)", value="three dogs playing chess, oil painting")
#openai.api_key = "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"
#number_of_images = 1
#image_response = openai.Image.create(prompt=command, n=number_of_images, size="512x512", response_format="url")
#for i in range(0, number_of_images):
#    url = image_response["data"][i]["url"]
#    response = requests.get(url)
#    img = Image.open(BytesIO(response.content))
#    st.image(img)
#--
st.subheader("Eksempel: Parametere")
with st.form("Input"):
    forfatter = st.text_input("Forfatter", value="Ola Nordmann")
    oppdragsleder = st.text_input("Oppdragsleder", value="Kari Nordmann")
    oppdragsgiver = st.text_input("Oppdragsgiver", value = "Firma AS")
    oppdragsnummer = st.text_input("Oppdragsnummer", value = "635960-01")
    sted = st.text_input("Sted", value = "Trondheim")
    #--
    depth_to_bedrock = st.number_input("Dybde til fjell [m]", value=5, step=1)
    loose_material = st.selectbox("Hva slags l√∏smasser?", options=["hav- og fjordavsetning", "elveavsetning", "breelvavsetning", "morene"])
    st.form_submit_button("Gi input")
#--
if depth_to_bedrock > 15:
    setning_dybde_til_fjell = "Siden dybde til fjell > 15 m, m√• det gj√∏res supplerende grunnunders√∏kelser. "
else:
    setning_dybde_til_fjell = ""
report_text_1 = f"""Det skal gj√∏res geoteknisk vurdering for {sted}. Dybde til fjell var {depth_to_bedrock} m. 
{setning_dybde_til_fjell}L√∏smassene er kartlagt som {loose_material}. Viktige problemstillinger innen geoteknikk er vurdering av fundamenters b√¶reevne (lastkapasitet) og deres setning under belastning, jordtrykk mot st√∏tte- og kjellermurer, stabilitet av veiskj√¶ringer og naturlige skr√•ninger, fundamentering av marine konstruksjoner og r√∏rledninger. """
#--
document = Document("src/data/docx/Termisk responstest.docx")
styles = document.styles
style = styles.add_style('Citation', WD_STYLE_TYPE.PARAGRAPH)
document.paragraphs[0].text = f"Oppdragsgiver: {oppdragsgiver}"
document.paragraphs[1].text = f"Oppdragsnavn: Geoteknisk rapport - {sted}"
document.paragraphs[2].text = f"Oppdragsnummer: {oppdragsnummer} - {sted}"
document.paragraphs[3].text = f"Utarbeidet av: {forfatter}"
document.paragraphs[4].text = f"Oppdragsleder: {oppdragsleder}"
document.paragraphs[5].text = f"Dato: {datetime.date.today()}"
document.paragraphs[6].text = f"Tilgjenglighet: √Öpen"

document.paragraphs[7].text = f"Geoteknisk notat - {sted}"

document.add_heading("Innledning", 1)
document.add_paragraph(report_text_1)

#--
st.markdown("---")
bio = io.BytesIO()
document.save(bio)
if document:
    st.download_button(
        label="Last ned rapport!",
        data=bio.getvalue(),
        file_name="Rapport.docx",
        mime="docx")
